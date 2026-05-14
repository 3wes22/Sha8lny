"""
Career Tools Service Layer

Handles resume building, ATS optimization, and portfolio management.
"""

from typing import Optional, List, Dict, Any
from decimal import Decimal
from django.db import transaction
from django.core.files.base import ContentFile

from .models import Resume, Portfolio
from apps.users.models import User
from apps.notifications.signals import resume_generated, resume_optimized, portfolio_published


class ResumeService:
    """Service for resume management and generation"""

    @staticmethod
    @transaction.atomic
    def create_resume(
        user: User,
        title: str,
        template_name: str = "modern",
        **resume_data
    ) -> Resume:
        """
        Create a new resume.

        Args:
            user: User instance
            title: Resume title
            template_name: Template to use
            **resume_data: Resume section data (personal_info, work_experience, etc.)

        Returns:
            Resume: Created resume
        """
        # If this is the first resume, make it primary
        existing_count = Resume.objects.filter(user=user, is_deleted=False).count()
        is_primary = existing_count == 0

        resume = Resume.objects.create(
            user=user,
            title=title,
            template_name=template_name,
            personal_info=resume_data.get('personal_info', {}),
            work_experience=resume_data.get('work_experience', {}),
            education=resume_data.get('education', {}),
            skills=resume_data.get('skills', {}),
            certifications=resume_data.get('certifications', {}),
            projects=resume_data.get('projects', {}),
            languages=resume_data.get('languages', {}),
            is_primary=is_primary,
            ats_processing_status='completed'
        )

        # Emit signal
        resume_generated.send(sender=Resume, instance=resume, user=user)

        return resume

    @staticmethod
    def get_user_resumes(user: User) -> List[Resume]:
        """Get all resumes for user"""
        return list(Resume.objects.filter(
            user=user,
            is_deleted=False
        ).order_by('-is_primary', '-updated_at'))

    @staticmethod
    def get_primary_resume(user: User) -> Optional[Resume]:
        """Get user's primary resume"""
        try:
            return Resume.objects.get(
                user=user,
                is_primary=True,
                is_deleted=False
            )
        except Resume.DoesNotExist:
            # Return most recent resume
            return Resume.objects.filter(
                user=user,
                is_deleted=False
            ).order_by('-updated_at').first()

    @staticmethod
    def get_resume_by_id(resume_id: str) -> Optional[Resume]:
        """Get resume by ID"""
        try:
            return Resume.objects.get(id=resume_id, is_deleted=False)
        except Resume.DoesNotExist:
            return None

    @staticmethod
    @transaction.atomic
    def update_resume(
        resume: Resume,
        **fields
    ) -> Resume:
        """Update resume data"""
        allowed_fields = [
            'title', 'template_name', 'personal_info', 'work_experience',
            'education', 'skills', 'certifications', 'projects', 'languages'
        ]

        for field, value in fields.items():
            if field in allowed_fields:
                setattr(resume, field, value)

        # Increment version
        resume.version += 1
        resume.save()

        return resume

    @staticmethod
    @transaction.atomic
    def set_primary_resume(resume: Resume) -> Resume:
        """Set resume as primary (unset others)"""
        # Unset all primary resumes for this user
        Resume.objects.filter(
            user=resume.user,
            is_primary=True,
            is_deleted=False
        ).update(is_primary=False)

        # Set this as primary
        resume.is_primary = True
        resume.save(update_fields=['is_primary', 'updated_at'])

        return resume

    @staticmethod
    def generate_pdf(resume: Resume) -> Resume:
        """
        Generate PDF file for resume using ReportLab.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm, mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        )
        from io import BytesIO

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=1.5 * cm,
            rightMargin=1.5 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )

        styles = getSampleStyleSheet()

        # Custom styles based on the user's template
        styles.add(ParagraphStyle(
            name='ResumeName',
            fontSize=26,
            leading=30,
            spaceAfter=4,
            textColor=HexColor('#1C2833'),
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name='ResumeTitle',
            fontSize=11,
            leading=14,
            spaceAfter=12,
            textColor=HexColor('#34495E'),
            fontName='Helvetica',
            alignment=TA_CENTER,
            textTransform='uppercase',
        ))
        styles.add(ParagraphStyle(
            name='ResumeContact',
            fontSize=10,
            leading=12,
            spaceBefore=6,
            spaceAfter=6,
            textColor=HexColor('#2C3E50'),
            fontName='Helvetica',
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name='ResumeSectionH2',
            fontSize=13,
            leading=16,
            spaceBefore=14,
            spaceAfter=2,
            textColor=HexColor('#1C2833'),
            fontName='Helvetica-Bold',
            textTransform='uppercase',
        ))
        styles.add(ParagraphStyle(
            name='ResumeBody',
            fontSize=10.5,
            leading=14,
            spaceAfter=4,
            textColor=HexColor('#000000'),
            fontName='Helvetica',
        ))
        styles.add(ParagraphStyle(
            name='ResumeBullet',
            fontSize=10.5,
            leading=14,
            spaceAfter=4,
            leftIndent=10,
            textColor=HexColor('#000000'),
            fontName='Helvetica',
        ))

        def section_heading(title):
            return [
                Paragraph(title.upper(), styles['ResumeSectionH2']),
                HRFlowable(width="100%", thickness=1, color=HexColor('#2C3E50'), spaceAfter=8, spaceBefore=2)
            ]

        elements = []

        # ---- HEADER ----
        personal = resume.personal_info or {}
        name = personal.get('name', personal.get('full_name', 'FULL NAME')).upper()
        elements.append(Paragraph(name, styles['ResumeName']))

        title = resume.title or "ENTRY LEVEL TARGET ROLE"
        elements.append(Paragraph(title.upper(), styles['ResumeTitle']))

        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#2C3E50')))
        
        contact_parts = []
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('linkedin'):
            contact_parts.append(personal['linkedin'])
        if personal.get('portfolio'):
            contact_parts.append(personal['portfolio'])
        elif personal.get('github'):
            contact_parts.append(personal['github'])
            
        if contact_parts:
            elements.append(Paragraph(' &nbsp;|&nbsp; '.join(contact_parts), styles['ResumeContact']))
            
        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#2C3E50')))

        # ---- OBJECTIVE / SUMMARY ----
        if personal.get('summary'):
            elements.extend(section_heading('OBJECTIVE'))
            elements.append(Paragraph(personal['summary'], styles['ResumeBody']))

        # ---- EDUCATION ----
        edu = resume.education
        edu_items = edu if isinstance(edu, list) else edu.get('items', []) if isinstance(edu, dict) else []
        if edu_items:
            elements.extend(section_heading('EDUCATION'))
            for ed in edu_items:
                # Degree | Major, University | Year (GPA)
                degree_parts = []
                if ed.get('degree'): degree_parts.append(ed['degree'])
                
                major_uni = []
                if ed.get('field'): major_uni.append(ed['field'])
                if ed.get('institution'): major_uni.append(ed['institution'])
                if major_uni: degree_parts.append(', '.join(major_uni))
                
                year_gpa = []
                if ed.get('graduation_date') or ed.get('dates'): 
                    year_gpa.append(ed.get('graduation_date', ed.get('dates')))
                if ed.get('gpa'):
                    year_gpa.append(f"(GPA: {ed['gpa']})")
                if year_gpa: degree_parts.append(' '.join(year_gpa))
                
                elements.append(Paragraph(' | '.join(degree_parts), styles['ResumeBody']))
                if ed.get('coursework'):
                    elements.append(Paragraph(f"Relevant Coursework: {ed['coursework']}", styles['ResumeBody']))
                elements.append(Spacer(1, 4))

        # ---- EXPERIENCE / INTERNSHIP ----
        work = resume.work_experience
        work_items = work if isinstance(work, list) else work.get('items', []) if isinstance(work, dict) else []
        if work_items:
            elements.extend(section_heading('EXPERIENCE'))
            for job in work_items:
                # Company | Intern | City, Country | MMM YYYY – MMM YYYY
                header_parts = []
                if job.get('company'): header_parts.append(job['company'])
                if job.get('role') or job.get('title'): header_parts.append(job.get('role', job.get('title')))
                if job.get('location'): header_parts.append(job['location'])
                
                dates = f"{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
                if dates.strip(' –'):
                    header_parts.append(dates)
                    
                elements.append(Paragraph(' | '.join(header_parts), styles['ResumeBody']))
                
                # Achievements
                achievements = job.get('achievements', job.get('description', []))
                if isinstance(achievements, str):
                    achievements = [achievements]
                for a in achievements:
                    elements.append(Paragraph(f'• {a}', styles['ResumeBullet']))
                elements.append(Spacer(1, 4))

        # ---- SKILLS ----
        skills = resume.skills
        skill_items = skills if isinstance(skills, list) else skills.get('items', []) if isinstance(skills, dict) else []
        if skill_items:
            elements.extend(section_heading('SKILLS'))
            # Try to group skills or just list them
            grouped_skills = {}
            flat_skills = []
            
            for s in skill_items:
                if isinstance(s, str):
                    flat_skills.append(s)
                elif isinstance(s, dict):
                    category = s.get('category', s.get('type'))
                    if category:
                        if category not in grouped_skills:
                            grouped_skills[category] = []
                        grouped_skills[category].append(s.get('name', ''))
                    else:
                        flat_skills.append(s.get('name', ''))
            
            if grouped_skills:
                for cat, items in grouped_skills.items():
                    elements.append(Paragraph(f"{cat.capitalize()}: {', '.join(items)}", styles['ResumeBody']))
                if flat_skills:
                    elements.append(Paragraph(f"Other: {', '.join(flat_skills)}", styles['ResumeBody']))
            elif flat_skills:
                elements.append(Paragraph(', '.join(flat_skills), styles['ResumeBody']))
            
            elements.append(Spacer(1, 4))

        # ---- CERTIFICATIONS ----
        certs = resume.certifications
        cert_items = certs if isinstance(certs, list) else certs.get('items', []) if isinstance(certs, dict) else []
        if cert_items:
            elements.extend(section_heading('CERTIFICATIONS'))
            for cert in cert_items:
                # Certification | Issuer | Year
                c_parts = []
                if cert.get('name'): c_parts.append(cert['name'])
                if cert.get('issuer'): c_parts.append(cert['issuer'])
                if cert.get('date') or cert.get('year'): c_parts.append(cert.get('date', cert.get('year')))
                elements.append(Paragraph(' | '.join(c_parts), styles['ResumeBody']))
            elements.append(Spacer(1, 4))

        # ---- PROJECTS ----
        projects = resume.projects
        project_items = projects if isinstance(projects, list) else projects.get('items', []) if isinstance(projects, dict) else []
        if project_items:
            elements.extend(section_heading('PROJECTS'))
            for proj in project_items:
                p_parts = []
                if proj.get('title'): p_parts.append(proj['title'])
                if proj.get('technologies'): 
                    techs = proj['technologies']
                    if isinstance(techs, list): techs = ', '.join(techs)
                    p_parts.append(techs)
                
                elements.append(Paragraph(' | '.join(p_parts), styles['ResumeBody']))
                
                if proj.get('description'):
                    desc = proj['description']
                    if isinstance(desc, list):
                        for d in desc: elements.append(Paragraph(f'• {d}', styles['ResumeBullet']))
                    else:
                        elements.append(Paragraph(f'• {desc}', styles['ResumeBullet']))
                elements.append(Spacer(1, 4))

        # ---- LANGUAGES ----
        langs = resume.languages
        lang_items = langs if isinstance(langs, list) else langs.get('items', []) if isinstance(langs, dict) else []
        if lang_items:
            elements.extend(section_heading('LANGUAGES'))
            lang_strs = []
            for lang in lang_items:
                if isinstance(lang, str):
                    lang_strs.append(lang)
                elif isinstance(lang, dict):
                    name = lang.get('name', lang.get('language', ''))
                    level = lang.get('proficiency', lang.get('level', ''))
                    lang_strs.append(f"{name} ({level})" if level else name)
            elements.append(Paragraph(', '.join(lang_strs), styles['ResumeBody']))

        # Build PDF
        doc.build(elements)

        # Save to model
        pdf_content = buffer.getvalue()
        buffer.close()

        resume.pdf_file.save(
            f"resume_{resume.id}.pdf",
            ContentFile(pdf_content),
            save=True
        )

        return resume

    @staticmethod
    def generate_docx(resume: Resume) -> Resume:
        """
        Generate DOCX file for resume using python-docx.
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # ---- HEADER ----
        personal = resume.personal_info or {}
        name = personal.get('name', personal.get('full_name', 'Your Name'))

        heading = doc.add_heading(name, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if personal.get('summary'):
            doc.add_paragraph(personal['summary'])

        # Helper to add section
        def add_section(title):
            doc.add_heading(title, level=1)

        # ---- WORK EXPERIENCE ----
        work = resume.work_experience
        work_items = work if isinstance(work, list) else work.get('items', []) if isinstance(work, dict) else []
        if work_items:
            add_section('Work Experience')
            for job in work_items:
                role = job.get('role', job.get('title', ''))
                company = job.get('company', '')
                p = doc.add_paragraph()
                run = p.add_run(f'{role}')
                run.bold = True
                if company:
                    p.add_run(f' — {company}')
                dates = f"{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
                date_para = doc.add_paragraph(dates)
                for run in date_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                achievements = job.get('achievements', job.get('description', []))
                if isinstance(achievements, str):
                    achievements = [achievements]
                for a in achievements:
                    doc.add_paragraph(a, style='List Bullet')

        # ---- EDUCATION ----
        edu = resume.education
        edu_items = edu if isinstance(edu, list) else edu.get('items', []) if isinstance(edu, dict) else []
        if edu_items:
            add_section('Education')
            for ed in edu_items:
                degree = ed.get('degree', '')
                field = ed.get('field', '')
                institution = ed.get('institution', '')
                line = degree
                if field:
                    line += f' in {field}'
                if institution:
                    line += f' — {institution}'
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                grad = ed.get('graduation_date', ed.get('dates', ''))
                if grad:
                    doc.add_paragraph(grad)

        # ---- SKILLS ----
        skills = resume.skills
        skill_items = skills if isinstance(skills, list) else skills.get('items', []) if isinstance(skills, dict) else []
        if skill_items:
            add_section('Skills')
            skill_names = []
            for s in skill_items:
                if isinstance(s, str):
                    skill_names.append(s)
                elif isinstance(s, dict):
                    skill_names.append(s.get('name', str(s)))
            doc.add_paragraph(', '.join(skill_names))

        # ---- PROJECTS ----
        projects = resume.projects
        project_items = projects if isinstance(projects, list) else projects.get('items', []) if isinstance(projects, dict) else []
        if project_items:
            add_section('Projects')
            for proj in project_items:
                p = doc.add_paragraph()
                run = p.add_run(proj.get('title', ''))
                run.bold = True
                if proj.get('description'):
                    doc.add_paragraph(proj['description'])

        # ---- CERTIFICATIONS ----
        certs = resume.certifications
        cert_items = certs if isinstance(certs, list) else certs.get('items', []) if isinstance(certs, dict) else []
        if cert_items:
            add_section('Certifications')
            for cert in cert_items:
                name = cert.get('name', '')
                issuer = cert.get('issuer', '')
                date = cert.get('date', '')
                line = name
                if issuer:
                    line += f' — {issuer}'
                if date:
                    line += f' ({date})'
                doc.add_paragraph(line)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()

        resume.docx_file.save(
            f"resume_{resume.id}.docx",
            ContentFile(docx_content),
            save=True
        )

        return resume

    @staticmethod
    @transaction.atomic
    def optimize_for_ats(resume: Resume) -> Resume:
        """
        Optimize resume for ATS systems.

        Tries Gemma via Ollama first, falls back to rule-based scoring.
        """
        import json
        import logging

        logger = logging.getLogger(__name__)

        # Mark as processing
        resume.ats_processing_status = 'processing'
        resume.save(update_fields=['ats_processing_status', 'updated_at'])

        # Build the resume text for analysis
        resume_text = ResumeService._build_resume_text(resume)

        try:
            result = ResumeService._ats_via_gemma(resume_text)
        except Exception as e:
            logger.warning(f"Gemma ATS failed, using fallback: {e}")
            result = ResumeService._ats_rule_based(resume)

        resume.ats_score = Decimal(str(result['score']))
        resume.is_ats_optimized = True
        resume.ats_suggestions = result['suggestions']
        resume.ats_processing_status = 'completed'
        resume.save()

        # Emit signal
        resume_optimized.send(sender=Resume, instance=resume)

        return resume

    @staticmethod
    def _build_resume_text(resume: Resume) -> str:
        """Flatten resume JSON into plain text for LLM analysis."""
        parts = []
        personal = resume.personal_info or {}
        if personal.get('name'):
            parts.append(f"Name: {personal['name']}")
        if personal.get('summary'):
            parts.append(f"Summary: {personal['summary']}")

        work = resume.work_experience
        work_items = work if isinstance(work, list) else work.get('items', []) if isinstance(work, dict) else []
        for job in work_items:
            parts.append(f"Role: {job.get('role', '')} at {job.get('company', '')}")
            for a in job.get('achievements', []):
                parts.append(f"  - {a}")

        skills = resume.skills
        skill_items = skills if isinstance(skills, list) else skills.get('items', []) if isinstance(skills, dict) else []
        skill_names = [s if isinstance(s, str) else s.get('name', '') for s in skill_items]
        if skill_names:
            parts.append(f"Skills: {', '.join(skill_names)}")

        return '\n'.join(parts)

    @staticmethod
    def _ats_via_gemma(resume_text: str) -> dict:
        """Call Gemma via Ollama for ATS analysis."""
        import httpx
        from apps.core.ai_settings import OLLAMA_HOST, OLLAMA_MODEL, OLLAMA_TIMEOUT

        prompt = f"""You are an ATS (Applicant Tracking System) optimization expert.

Analyze this resume and return a JSON object with exactly two keys:
- "score": an integer from 0-100 rating ATS-friendliness
- "suggestions": an array of 5-8 specific, actionable improvement strings

Resume:
---
{resume_text}
---

Rules for scoring:
- 90-100: Strong action verbs, quantified achievements, standard section headings, keyword-rich
- 70-89: Good structure but missing some keywords or quantification
- 50-69: Weak formatting, vague descriptions, missing sections
- Below 50: Major issues (no clear structure, no keywords)

Return ONLY valid JSON, no markdown, no explanation."""

        response = httpx.post(
            f"{OLLAMA_HOST}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.2, "num_ctx": 4096},
            },
            timeout=OLLAMA_TIMEOUT,
        )
        response.raise_for_status()

        import json
        raw = response.json()["response"]

        # Try to parse JSON from the response
        # Sometimes Gemma wraps in ```json...```
        clean = raw.strip()
        if clean.startswith('```'):
            clean = clean.split('\n', 1)[1] if '\n' in clean else clean[3:]
            if clean.endswith('```'):
                clean = clean[:-3]
            clean = clean.strip()

        result = json.loads(clean)
        score = max(0, min(100, int(result.get('score', 50))))
        suggestions = result.get('suggestions', [])

        return {'score': score, 'suggestions': {'improvements': suggestions}}

    @staticmethod
    def _ats_rule_based(resume: Resume) -> dict:
        """Deterministic fallback when Gemma is unavailable."""
        score = 30  # base
        suggestions = []

        personal = resume.personal_info or {}
        if personal.get('name'):
            score += 5
        else:
            suggestions.append('Add your full name to personal info')
        if personal.get('email'):
            score += 5
        else:
            suggestions.append('Add a professional email address')
        if personal.get('phone'):
            score += 3
        if personal.get('summary'):
            score += 7
            if len(personal['summary']) < 50:
                suggestions.append('Expand your professional summary to at least 2-3 sentences')
        else:
            suggestions.append('Add a professional summary highlighting your key qualifications')

        work = resume.work_experience
        work_items = work if isinstance(work, list) else work.get('items', []) if isinstance(work, dict) else []
        if work_items:
            score += 10
            for job in work_items:
                achieves = job.get('achievements', [])
                if len(achieves) >= 2:
                    score += 3
        else:
            suggestions.append('Add work experience with specific achievements')

        edu = resume.education
        edu_items = edu if isinstance(edu, list) else edu.get('items', []) if isinstance(edu, dict) else []
        if edu_items:
            score += 8
        else:
            suggestions.append('Add your education background')

        skills = resume.skills
        skill_items = skills if isinstance(skills, list) else skills.get('items', []) if isinstance(skills, dict) else []
        if len(skill_items) >= 5:
            score += 10
        elif skill_items:
            score += 5
            suggestions.append('Add more relevant skills — aim for at least 5-8 key skills')
        else:
            suggestions.append('Add a skills section with relevant technical and soft skills')

        if resume.projects and resume.projects != {}:
            score += 5
        else:
            suggestions.append('Add projects to demonstrate hands-on experience')

        if resume.certifications and resume.certifications != {}:
            score += 5

        if not suggestions:
            suggestions.append('Your resume looks well-structured. Consider adding more quantified achievements.')

        score = min(100, score)

        return {
            'score': score,
            'suggestions': {'improvements': suggestions},
        }


class PortfolioService:
    """Service for portfolio management"""

    @staticmethod
    @transaction.atomic
    def create_portfolio(
        user: User,
        title: str,
        description: str = "",
        custom_url_slug: Optional[str] = None,
        theme: str = "default"
    ) -> Portfolio:
        """
        Create a new portfolio.

        Args:
            user: User instance
            title: Portfolio title
            description: Portfolio description
            custom_url_slug: Custom URL slug
            theme: Theme to use

        Returns:
            Portfolio: Created portfolio
        """
        portfolio = Portfolio.objects.create(
            user=user,
            title=title,
            description=description,
            custom_url_slug=custom_url_slug,
            theme=theme,
            is_public=False  # Start as private
        )

        return portfolio

    @staticmethod
    def get_user_portfolios(user: User) -> List[Portfolio]:
        """Get all portfolios for user"""
        return list(Portfolio.objects.filter(
            user=user,
            is_deleted=False
        ).order_by('-updated_at'))

    @staticmethod
    def get_portfolio_by_id(portfolio_id: str) -> Optional[Portfolio]:
        """Get portfolio by ID"""
        try:
            return Portfolio.objects.get(id=portfolio_id, is_deleted=False)
        except Portfolio.DoesNotExist:
            return None

    @staticmethod
    def get_portfolio_by_slug(slug: str) -> Optional[Portfolio]:
        """Get public portfolio by custom URL slug"""
        try:
            return Portfolio.objects.get(
                custom_url_slug=slug,
                is_public=True,
                is_deleted=False
            )
        except Portfolio.DoesNotExist:
            return None

    @staticmethod
    @transaction.atomic
    def update_portfolio(
        portfolio: Portfolio,
        **fields
    ) -> Portfolio:
        """Update portfolio data"""
        allowed_fields = [
            'title', 'description', 'projects', 'achievements',
            'testimonials', 'theme', 'custom_styles', 'custom_url_slug'
        ]

        for field, value in fields.items():
            if field in allowed_fields:
                setattr(portfolio, field, value)

        portfolio.save()
        return portfolio

    @staticmethod
    @transaction.atomic
    def publish_portfolio(portfolio: Portfolio) -> Portfolio:
        """Make portfolio public"""
        portfolio.is_public = True
        portfolio.save(update_fields=['is_public', 'updated_at'])

        # Emit signal
        portfolio_published.send(sender=Portfolio, instance=portfolio)

        return portfolio

    @staticmethod
    @transaction.atomic
    def unpublish_portfolio(portfolio: Portfolio) -> Portfolio:
        """Make portfolio private"""
        portfolio.is_public = False
        portfolio.save(update_fields=['is_public', 'updated_at'])

        return portfolio

    @staticmethod
    def increment_view_count(portfolio: Portfolio) -> None:
        """Increment portfolio view count"""
        portfolio.increment_view_count()

    @staticmethod
    @transaction.atomic
    def add_project_to_portfolio(
        portfolio: Portfolio,
        project_data: Dict[str, Any]
    ) -> Portfolio:
        """
        Add project to portfolio.

        Args:
            portfolio: Portfolio instance
            project_data: Project information dict

        Returns:
            Portfolio: Updated portfolio
        """
        projects = portfolio.projects

        # Initialize if needed
        if not isinstance(projects, dict):
            projects = {'items': []}
        if 'items' not in projects:
            projects['items'] = []

        # Add new project
        projects['items'].append(project_data)
        portfolio.projects = projects
        portfolio.save()

        return portfolio

    @staticmethod
    @transaction.atomic
    def add_achievement_to_portfolio(
        portfolio: Portfolio,
        achievement_data: Dict[str, Any]
    ) -> Portfolio:
        """Add achievement to portfolio"""
        achievements = portfolio.achievements

        if not isinstance(achievements, dict):
            achievements = {'items': []}
        if 'items' not in achievements:
            achievements['items'] = []

        achievements['items'].append(achievement_data)
        portfolio.achievements = achievements
        portfolio.save()

        return portfolio

    @staticmethod
    @transaction.atomic
    def add_testimonial_to_portfolio(
        portfolio: Portfolio,
        testimonial_data: Dict[str, Any]
    ) -> Portfolio:
        """Add testimonial to portfolio"""
        testimonials = portfolio.testimonials

        if not isinstance(testimonials, dict):
            testimonials = {'items': []}
        if 'items' not in testimonials:
            testimonials['items'] = []

        testimonials['items'].append(testimonial_data)
        portfolio.testimonials = testimonials
        portfolio.save()

        return portfolio
